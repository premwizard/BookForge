import docx
from docx.shared import Pt, Inches
import io

class FormattingEngine:
    def __init__(self, template_settings: dict = None):
        self.settings = template_settings or {
            "font": "Times New Roman",
            "size": 12,
            "margin": 1.0, # Inches
            "line_spacing": 1.5
        }

    def generate_docx(self, ai_metadata: dict, raw_text: str) -> bytes:
        """
        Creates a formatted DOCX file based on AI metadata and settings.
        Returns the bytes of the generated DOCX.
        """
        doc = docx.Document()
        
        # Apply margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(self.settings["margin"])
            section.bottom_margin = Inches(self.settings["margin"])
            section.left_margin = Inches(self.settings["margin"])
            section.right_margin = Inches(self.settings["margin"])

        # Write Title
        if "title" in ai_metadata:
            title = doc.add_heading(ai_metadata["title"], 0)
            title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Write Chapters (stub)
        if "chapters" in ai_metadata:
            for chapter in ai_metadata["chapters"]:
                doc.add_heading(chapter.get("title", "Chapter"), level=1)
                
                # We would normally match sections from raw_text here
                p = doc.add_paragraph("Chapter content would go here...")
                p.paragraph_format.line_spacing = self.settings["line_spacing"]
                
                for section in chapter.get("sections", []):
                    doc.add_heading(section, level=2)
                    doc.add_paragraph("Section content...")

        # Save to memory
        byte_io = io.BytesIO()
        doc.save(byte_io)
        return byte_io.getvalue()
