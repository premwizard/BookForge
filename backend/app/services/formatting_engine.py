import docx
import os
from uuid import UUID
from sqlalchemy.orm import Session
from app.models.parser import ParsedDocument, ParsedElement, DocumentTable, DocumentImage
from app.models.template import TemplateVersion, StyleMapping

class FormattingEngine:
    def __init__(self, db: Session, parsed_doc_id: UUID, template_version_id: UUID):
        self.db = db
        self.parsed_doc_id = parsed_doc_id
        self.template_version_id = template_version_id
        
        self.parsed_doc = db.query(ParsedDocument).filter(ParsedDocument.id == parsed_doc_id).first()
        self.template_version = db.query(TemplateVersion).filter(TemplateVersion.id == template_version_id).first()
        
        if not self.parsed_doc or not self.template_version:
            raise ValueError("Invalid ParsedDocument or TemplateVersion")
            
        # Load style mappings
        self.mappings = {}
        mappings_db = db.query(StyleMapping).filter(StyleMapping.template_version_id == template_version_id).all()
        for mapping in mappings_db:
            self.mappings[mapping.raw_element_type] = mapping.template_style.style_name
            
        # Initialize output document based on the base template
        self.doc = docx.Document(self.template_version.storage_path)
        # Clear existing paragraphs if any in the base template, but preserve styles
        for paragraph in self.doc.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None
            
    def _get_mapped_style(self, raw_element_type: str) -> str:
        """Get the mapped publisher style name, or fallback to Normal"""
        return self.mappings.get(raw_element_type, "Normal")

    def generate(self, output_path: str):
        """
        Traverse the document tree and generate the final formatted DOCX.
        """
        # Fetch root elements
        elements = self.db.query(ParsedElement).filter(
            ParsedElement.parsed_document_id == self.parsed_doc_id,
            ParsedElement.parent_id == None
        ).order_by(ParsedElement.position).all()
        
        for element in elements:
            self._process_element(element)
            
        # Save output
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        return output_path
        
    def _process_element(self, element: ParsedElement):
        """Process a single element and its children recursively."""
        style_name = self._get_mapped_style(element.element_type)
        
        try:
            if element.element_type == "Paragraph":
                self.doc.add_paragraph(element.text or "", style=style_name)
            elif element.element_type.startswith("Heading"):
                self.doc.add_paragraph(element.text or "", style=style_name)
            elif element.element_type == "Table":
                self._process_table(element, style_name)
            elif element.element_type == "Image":
                self._process_image(element, style_name)
            elif element.element_type == "PageBreak":
                self.doc.add_page_break()
            else:
                # Default text processing
                self.doc.add_paragraph(element.text or "", style=style_name)
        except KeyError:
            # Fallback if style doesn't exist in the document
            self.doc.add_paragraph(element.text or "", style="Normal")
            
        # Process children
        for child in element.children:
            self._process_element(child)

    def _process_table(self, element: ParsedElement, style_name: str):
        table_data = self.db.query(DocumentTable).filter(DocumentTable.element_id == element.id).first()
        if not table_data or not table_data.table_data:
            return
            
        grid = table_data.table_data.get("grid", [])
        if not grid:
            return
            
        rows = len(grid)
        cols = len(grid[0]) if rows > 0 else 0
        
        doc_table = self.doc.add_table(rows=rows, cols=cols)
        try:
            doc_table.style = style_name
        except KeyError:
            pass # Fallback to default table style
            
        for r_idx, row in enumerate(grid):
            for c_idx, cell_value in enumerate(row):
                if c_idx < cols:
                    doc_table.cell(r_idx, c_idx).text = str(cell_value)
                    
        if table_data.caption:
            self.doc.add_paragraph(table_data.caption, style=self._get_mapped_style("Caption"))

    def _process_image(self, element: ParsedElement, style_name: str):
        image_data = self.db.query(DocumentImage).filter(DocumentImage.element_id == element.id).first()
        if not image_data or not os.path.exists(image_data.storage_path):
            return
            
        self.doc.add_picture(image_data.storage_path) # optionally specify width from attributes
        if image_data.caption:
            self.doc.add_paragraph(image_data.caption, style=self._get_mapped_style("Caption"))
