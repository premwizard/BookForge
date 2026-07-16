from app.document_parser.base import BaseParser, ParsedDocumentDTO, ParsedElementDTO
import docx

class DocxParser(BaseParser):
    def parse(self, file_path: str, original_filename: str = None) -> ParsedDocumentDTO:
        doc = docx.Document(file_path)
        parsed_doc = ParsedDocumentDTO(metadata={"filename": original_filename, "source": "docx"})
        
        # Extract Core Properties
        core_props = doc.core_properties
        parsed_doc.metadata.update({
            "title": core_props.title,
            "author": core_props.author,
            "creation_date": str(core_props.created) if core_props.created else None
        })

        # Process Paragraphs
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
                
            element_type = "Paragraph"
            if p.style.name.startswith("Heading"):
                element_type = p.style.name.replace(" ", "")
                
            el = ParsedElementDTO(
                element_type=element_type,
                text=p.text,
                attributes={"style": p.style.name, "alignment": str(p.alignment)}
            )
            parsed_doc.add_element(el)
            
        # Note: In a production scenario, we'd iterate through document.element.body.iter() 
        # to preserve the exact order of paragraphs, tables, and images.
        
        # Process Tables
        for table in doc.tables:
            table_el = ParsedElementDTO(element_type="Table", attributes={"style": table.style.name})
            for row in table.rows:
                row_el = ParsedElementDTO(element_type="TableRow")
                for cell in row.cells:
                    cell_el = ParsedElementDTO(element_type="TableCell", text=cell.text.strip())
                    row_el.add_child(cell_el)
                table_el.add_child(row_el)
            parsed_doc.add_element(table_el)

        return parsed_doc
